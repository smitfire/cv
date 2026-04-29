#!/usr/bin/env python3
"""
render.py — Build CV (PDF + DOCX) and cover letter (PDF) for one application.

Usage:
    python scripts/render.py applications/<slug>/

Reads:
    <slug>/tailored.yaml    — self-contained data for this application
    templates/cv.html.j2
    templates/cover_letter.md.j2

Writes:
    <slug>/cv.pdf           — ATS-safe single-column PDF, with PDF metadata
    <slug>/cv.docx          — ATS-friendly DOCX, with core properties
    <slug>/cover_letter.pdf — Markdown → HTML → PDF
"""
from __future__ import annotations

import argparse
import re
import sys
from datetime import date, datetime
from pathlib import Path

import yaml
from jinja2 import Environment, FileSystemLoader, StrictUndefined
from markupsafe import Markup

ROOT = Path(__file__).resolve().parent.parent
TEMPLATES = ROOT / "templates"

# Graduated cap on bullets per role — older roles get fewer, since the most
# recent / most-relevant role should dominate the page.
BULLET_CAP_BY_ROLE_INDEX = [5, 4, 3, 3, 2, 2, 2, 2]
MAX_BULLETS_PER_ROLE = 5  # absolute fallback


# ---------------------------------------------------------------------------- helpers


def _fmt_date(v) -> str:
    """Render dates in human form: 'Jul 2025' / '2025' / 'Present'."""
    if v is None:
        return ""
    if isinstance(v, (date, datetime)):
        return v.strftime("%b %Y")
    if isinstance(v, int):
        return str(v)
    if isinstance(v, str):
        # Try YYYY-MM
        m = re.match(r"^(\d{4})-(\d{1,2})$", v)
        if m:
            year, month = int(m.group(1)), int(m.group(2))
            try:
                return datetime(year, month, 1).strftime("%b %Y")
            except ValueError:
                pass
        # YYYY-MM-DD
        m = re.match(r"^(\d{4})-(\d{1,2})-\d{1,2}$", v)
        if m:
            year, month = int(m.group(1)), int(m.group(2))
            try:
                return datetime(year, month, 1).strftime("%b %Y")
            except ValueError:
                pass
        # Plain YYYY
        if re.match(r"^\d{4}$", v):
            return v
        return v
    return str(v)


# Bold common metric tokens in body text. Applied as a Jinja filter; output is HTML-safe.
import html as _html

_METRIC_PATTERNS = [
    re.compile(r"(\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*[×xX]\b)"),    # 3,370× / 10x
    re.compile(r"(\b\d+(?:\.\d+)?\s*%)"),                          # 5% / 85%
    re.compile(r"(\b\d+(?:\.\d+)?\s*(?:to|→|->)\s*\d+(?:\.\d+)?\s*%)"),  # 5% to 85%
    re.compile(r"(\b\d{2,}\+\b)"),                                # 320+, 10k+
    re.compile(r"(\b\d{1,3}(?:,\d{3})+\b)"),                      # 162,000
    re.compile(r"(\$\s*\d+(?:\.\d+)?\s*[KMB]?\b)"),               # $50M
    re.compile(r"(\b€\s*\d+(?:\.\d+)?\s*[KMB]?\b)"),              # €4.4M
]


def metric_emphasis(text: str) -> Markup:
    """HTML-escape input, then wrap metric tokens in <strong>. Markup-safe (won't double-escape)."""
    if text is None:
        return Markup("")
    safe = _html.escape(str(text))
    for pat in _METRIC_PATTERNS:
        safe = pat.sub(r"<strong>\1</strong>", safe)
    return Markup(safe)


def _normalise_experience(exp_list):
    """Add start_str/end_str fields the template expects, plus default empty fields
    so StrictUndefined doesn't blow up on optional keys. Graduated cap on bullets
    per role — recent roles get more, older roles get fewer."""
    out = []
    for i, e in enumerate(exp_list):
        e = dict(e)
        e["start_str"] = _fmt_date(e.get("start"))
        end = e.get("end")
        if end in (None, "present", "Present"):
            e["end_str"] = "Present"
        else:
            e["end_str"] = _fmt_date(end)
        e.setdefault("blurb", "")
        e.setdefault("stack", [])
        e.setdefault("bullets", [])
        e.setdefault("location", "")
        # Graduated bullet cap (most recent role dominates).
        cap = (
            BULLET_CAP_BY_ROLE_INDEX[i]
            if i < len(BULLET_CAP_BY_ROLE_INDEX)
            else BULLET_CAP_BY_ROLE_INDEX[-1]
        )
        if e["bullets"] and len(e["bullets"]) > cap:
            e["bullets"] = e["bullets"][:cap]
        out.append(e)
    return out


def _load_tailored(app_dir: Path) -> dict:
    path = app_dir / "tailored.yaml"
    if not path.exists():
        sys.exit(f"error: {path} not found. run scripts/new_application.py first.")
    with path.open() as f:
        return yaml.safe_load(f)


def _jinja(template: str):
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATES)),
        undefined=StrictUndefined,
        trim_blocks=False,
        lstrip_blocks=False,
        autoescape=True,
    )
    env.filters["metric_emphasis"] = metric_emphasis
    return env.get_template(template)


# ------------------------------------------------------------------------- pdf render


def _render_cv_pdf(data: dict, html_str: str, out_path: Path, meta: dict) -> None:
    from weasyprint import HTML

    HTML(string=html_str, base_url=str(ROOT)).write_pdf(str(out_path))
    _set_pdf_metadata(out_path, meta)


def _set_pdf_metadata(pdf_path: Path, meta: dict) -> None:
    """Set /Title /Author /Subject /Keywords on a rendered PDF."""
    try:
        import pikepdf
    except ImportError:
        print("warn: pikepdf not installed; skipping PDF metadata write", file=sys.stderr)
        return
    with pikepdf.open(str(pdf_path), allow_overwriting_input=True) as pdf:
        with pdf.open_metadata(set_pikepdf_as_editor=False) as md:
            if meta.get("title"):    md["dc:title"]       = meta["title"]
            if meta.get("author"):   md["dc:creator"]     = [meta["author"]]
            if meta.get("subject"):  md["dc:description"] = meta["subject"]
            if meta.get("keywords"):
                md["pdf:Keywords"] = ", ".join(meta["keywords"])
        # Ensure legacy info dict is also populated (older ATS read this).
        pdf.docinfo["/Title"]    = meta.get("title", "")
        pdf.docinfo["/Author"]   = meta.get("author", "")
        pdf.docinfo["/Subject"]  = meta.get("subject", "")
        if meta.get("keywords"):
            pdf.docinfo["/Keywords"] = ", ".join(meta["keywords"])
        pdf.save(str(pdf_path))


# ------------------------------------------------------------------------ docx render


def _render_cv_docx(data: dict, out_path: Path, meta: dict) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # name
    h = doc.add_paragraph()
    run = h.add_run(data["basics"]["name"])
    run.bold = True
    run.font.size = Pt(20)

    # role line
    p = doc.add_paragraph()
    r = p.add_run(data["headline"])
    r.font.size = Pt(12)

    # contact
    contact_bits = [
        data["basics"].get("location", ""),
        data["basics"].get("phone", ""),
        data["basics"].get("email", ""),
        data["basics"].get("linkedin", ""),
    ]
    sites = data["basics"].get("sites") or {}
    if sites.get("personal"):
        contact_bits.append(sites["personal"])
    p = doc.add_paragraph()
    r = p.add_run("  ·  ".join(b for b in contact_bits if b))
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def section(title: str):
        h = doc.add_paragraph()
        r = h.add_run(title.upper())
        r.bold = True
        r.font.size = Pt(11)

    # summary
    section("Summary")
    doc.add_paragraph(data["summary"])

    # experience
    section("Experience")
    for role in data["experience"]:
        h = doc.add_paragraph()
        r = h.add_run(f"{role['role']} — {role['company']}")
        r.bold = True
        start_str = role.get("start_str") or _fmt_date(role.get("start"))
        end = role.get("end")
        end_str = "Present" if end in (None, "present", "Present") else (role.get("end_str") or _fmt_date(end))
        loc = f" · {role['location']}" if role.get("location") else ""
        r2 = h.add_run(f"   ({start_str} – {end_str}{loc})")
        r2.font.size = Pt(10)
        r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        if role.get("blurb"):
            p = doc.add_paragraph(role["blurb"])
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(10)

        for b in role.get("bullets", []):
            doc.add_paragraph(b, style="List Bullet")

        if role.get("stack"):
            p = doc.add_paragraph()
            r = p.add_run("Stack: ")
            r.bold = True
            p.add_run(", ".join(role["stack"]))

    # achievements
    if data.get("key_achievements"):
        section("Selected Achievements")
        for a in data["key_achievements"]:
            doc.add_paragraph(a, style="List Bullet")

    # skills
    section("Skills")
    for s in data["skills"]:
        p = doc.add_paragraph()
        r = p.add_run(f"{s['label']}: ")
        r.bold = True
        p.add_run(", ".join(s["items"]))

    # projects
    if data.get("projects"):
        section("Selected Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{proj['title']} — ")
            r.bold = True
            p.add_run(proj["blurb"])

    # education
    section("Education")
    for e in data["education"]:
        loc = f", {e['location']}" if e.get("location") else ""
        end = f" – {e['end']}" if e.get("end") else ""
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{e['school']} ")
        r.bold = True
        p.add_run(f"— {e['degree']}{loc} ({e['start']}{end})")

    # languages
    section("Languages")
    for l in data["languages"]:
        doc.add_paragraph(f"{l['name']} — {l['level']}", style="List Bullet")

    # core properties (DOCX metadata; some ATS read this).
    # Note: docx `keywords` core property has a 255-char limit — truncate by dropping
    # tail keywords until we fit, instead of raising.
    cp = doc.core_properties
    cp.title    = (meta.get("title") or "")[:255]
    cp.author   = (meta.get("author") or "")[:255]
    cp.subject  = (meta.get("subject") or "")[:255]
    if meta.get("keywords"):
        kws = list(meta["keywords"])
        joined = ", ".join(kws)
        while len(joined) > 255 and len(kws) > 1:
            kws.pop()
            joined = ", ".join(kws)
        cp.keywords = joined[:255]

    doc.save(str(out_path))


# ----------------------------------------------------------------- cover letter render


def _render_cover_letter(data: dict, out_path: Path, meta: dict) -> None:
    cl = data.get("cover_letter")
    if not cl:
        print("note: no cover_letter section in tailored.yaml — skipping", file=sys.stderr)
        return

    tpl = _jinja("cover_letter.md.j2")
    md_str = tpl.render(
        basics=data["basics"],
        today=cl.get("date") or date.today().strftime("%-d %B %Y"),
        hiring_manager=cl.get("hiring_manager", ""),
        company=cl["company"],
        company_location=cl.get("company_location", ""),
        role_title=cl["role_title"],
        salutation=cl.get("salutation", "Hiring Team"),
        opening=cl["opening"],
        body=cl["body"],
        closing=cl.get("closing"),
    )

    import markdown as md_lib
    body_html = md_lib.markdown(md_str)
    html_str = f"""<!DOCTYPE html><html><head><meta charset='UTF-8'>
<style>
@page {{ size: A4; margin: 22mm 22mm; }}
html, body {{ font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif;
  font-size: 11pt; line-height: 1.5; color: #111; }}
p {{ margin: 0 0 10pt 0; }}
strong {{ color: #111; }}
a {{ color: #2c5282; text-decoration: none; }}
</style></head><body>{body_html}</body></html>"""

    from weasyprint import HTML
    HTML(string=html_str, base_url=str(ROOT)).write_pdf(str(out_path))
    _set_pdf_metadata(out_path, meta)


# ---------------------------------------------------------------------------- driver


def render(app_dir: Path) -> None:
    data = _load_tailored(app_dir)
    data["experience"] = _normalise_experience(data["experience"])

    cv_meta = {
        "title":   f"{data['basics']['name']} — {data['headline']}",
        "author":  data["basics"]["name"],
        "subject": f"CV for {data.get('meta', {}).get('role_title', '')} at {data.get('meta', {}).get('company', '')}".strip(),
        "keywords": data.get("meta", {}).get("jd_keywords", []),
    }
    cl_meta = {
        **cv_meta,
        "title": f"{data['basics']['name']} — Cover Letter ({data.get('meta', {}).get('company', '')})",
        "subject": f"Cover letter for {data.get('meta', {}).get('role_title', '')} at {data.get('meta', {}).get('company', '')}".strip(),
    }

    # CV → HTML → PDF
    tpl = _jinja("cv.html.j2")
    html_str = tpl.render(
        basics=data["basics"],
        headline=data["headline"],
        summary=data["summary"],
        experience=data["experience"],
        skills=data["skills"],
        education=data["education"],
        languages=data["languages"],
        projects=data.get("projects", []),
        key_achievements=data.get("key_achievements", []),
    )
    (app_dir / "cv.html").write_text(html_str)
    print(f"  wrote {app_dir / 'cv.html'}")

    pdf_path = app_dir / "cv.pdf"
    _render_cv_pdf(data, html_str, pdf_path, cv_meta)
    print(f"  wrote {pdf_path}")

    docx_path = app_dir / "cv.docx"
    _render_cv_docx(data, docx_path, cv_meta)
    print(f"  wrote {docx_path}")

    cl_path = app_dir / "cover_letter.pdf"
    _render_cover_letter(data, cl_path, cl_meta)
    if cl_path.exists():
        print(f"  wrote {cl_path}")


def main() -> None:
    p = argparse.ArgumentParser()
    p.add_argument("app_dir", help="path to applications/<slug>/")
    args = p.parse_args()

    app_dir = Path(args.app_dir).resolve()
    if not app_dir.is_dir():
        sys.exit(f"error: {app_dir} is not a directory")
    print(f"rendering {app_dir.name}")
    render(app_dir)


if __name__ == "__main__":
    main()
